"""Export each publication CSV table in tables/ as a formatted .docx table.

Wide tables (model-comparison CSVs carry 11-12 columns of diagnostic detail:
R2, RMSE, MAE, timing, both speedup ratios, CV mean/std, split sizes, ...)
do not fit a portrait Word page at a readable font size -- the same
overflow problem the PDF had. For those specifically we (a) drop the
redundant/secondary columns (MAE, Speedup_vs_ANN, Variant/N_train/N_test --
still in the CSV for anyone who wants them), (b) shorten the model names,
and (c) render the page in landscape with narrow margins, which between them
keep every remaining column comfortably inside the printable width.
"""
import glob
import os
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
TABLES = os.path.join(ROOT, "tables")

TITLES = {
    "dataset_summary": "Table 1. Summary of the assembled flexible and rigid LTPP IRI datasets",
    "master_lightgbm_comparison": "Table 2. LightGBM held-out R² across every pavement family and model variant",
    "flexible_structural_model_metrics": "Table 3. Structural model comparison — flexible (asphalt) pavements",
    "rigid_structural_model_metrics": "Table 4. Structural model comparison — rigid pavements, pooled",
}

SHORT_MODEL_NAMES = {
    "MEPDG-style Linear Regression": "Linear (M-E style)",
    "Random Forest": "Random Forest",
    "ANN (MLP, Optuna-tuned)": "ANN (tuned)",
    "Gradient-Boosted Ensemble (LightGBM, Optuna-tuned)": "LightGBM (tuned)",
}

# Columns kept for the wide *_model_metrics tables (rest stay in the CSV).
METRICS_DISPLAY_COLS = ["Model", "R2", "RMSE", "Train_time_s", "Speedup_vs_RF",
                         "CV_R2_mean", "CV_R2_std"]
LANDSCAPE_PATTERNS = ("_model_metrics", "master_lightgbm_comparison", "all_variants_metrics")


def export_docx(csv_path):
    name = os.path.splitext(os.path.basename(csv_path))[0]
    df = pd.read_csv(csv_path)
    if df.columns[0].startswith("Unnamed"):
        df = df.rename(columns={df.columns[0]: "Feature"})

    is_wide_metrics = "_model_metrics" in name
    if is_wide_metrics:
        if "Model" in df.columns:
            df["Model"] = df["Model"].replace(SHORT_MODEL_NAMES)
        keep = [c for c in METRICS_DISPLAY_COLS if c in df.columns]
        df = df[keep]

    df = df.round(4)

    doc = Document()
    if any(p in name for p in LANDSCAPE_PATTERNS):
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
        section.left_margin = section.right_margin = Inches(0.6)
        section.top_margin = section.bottom_margin = Inches(0.6)

    title = TITLES.get(name, name)
    h = doc.add_heading(title, level=2)
    for run in h.runs:
        run.font.size = Pt(12)
    if is_wide_metrics:
        note = doc.add_paragraph(
            "Additional columns (MAE, Speedup_vs_ANN, Variant, N_train, N_test) are in the "
            f"source CSV: tables/{name}.csv")
        note.runs[0].italic = True
        note.runs[0].font.size = Pt(8)

    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Light Grid Accent 1"
    table.autofit = True
    hdr = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr[i].text = str(col)
        for p in hdr[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.bold = True

    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = "" if pd.isna(val) else str(val)

    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)

    out_path = os.path.join(TABLES, f"{name}.docx")
    doc.save(out_path)
    print(f"wrote {out_path}")


if __name__ == "__main__":
    for csv_path in glob.glob(os.path.join(TABLES, "*.csv")):
        export_docx(csv_path)
